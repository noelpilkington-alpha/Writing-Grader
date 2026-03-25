"""Generate a formatted DOCX coaching report for Zara Lapointe G3.2."""

import os
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding="utf-8")

doc = Document()

# -- Style setup --
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)


def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def bold(para, text):
    run = para.add_run(text)
    run.bold = True
    return run


def italic(para, text):
    run = para.add_run(text)
    run.italic = True
    return run


def bold_italic(para, text):
    run = para.add_run(text)
    run.bold = True
    run.italic = True
    return run


def shade_cell(cell, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def table_header(table, texts, color="E8EAF6"):
    for i, text in enumerate(texts):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        shade_cell(cell, color)


def table_add(table, cells):
    row = table.add_row()
    for i, text in enumerate(cells):
        c = row.cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(text)
        run.font.size = Pt(10)
    return row


# =====================================================================
# TITLE
# =====================================================================
title = doc.add_heading("Writing Coach Report: Zara Lapointe \u2014 G3.2", level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

# =====================================================================
# WHAT YOU DID WELL
# =====================================================================
add_heading_styled("What You Did Well", level=1)

p = doc.add_paragraph(
    "Zara, you showed a lot of good thinking on this test! "
    "Here\u2019s what you should be proud of:"
)

# Q3
p = doc.add_paragraph(style="List Bullet")
bold(p, "Q3")
p.add_run(" \u2014 You turned a statement into a question perfectly: ")
italic(p, "\u201cDoes this fun event sell out every year?\u201d")
p.add_run(" That shows you know how to change sentence types.")

# Q2
p = doc.add_paragraph(style="List Bullet")
bold(p, "Q2")
p.add_run(
    " \u2014 You combined two sentences using \u201cbecause\u201d "
    "and the meaning was clear. Nice work!"
)

# Q8
p = doc.add_paragraph(style="List Bullet")
bold(p, "Q8")
p.add_run(" \u2014 ")
italic(
    p,
    "\u201cYes, I would love to attend a Teddy Bear Toss game because "
    "it sounds fun, and I also love Teddy Bears.\u201d",
)
p.add_run(
    " \u2014 This was your best one-sentence response. You answered the "
    "question, gave two reasons, and kept it all in one complete sentence."
)

# Reasons
p = doc.add_paragraph(style="List Bullet")
bold(p, "You try to give reasons")
p.add_run(
    " \u2014 That\u2019s a great habit! Most of your Q6\u2013Q9 responses "
    "include \u201cbecause\u201d or an explanation, which is exactly what "
    "the rubric is looking for."
)

# =====================================================================
# THE 3 THINGS
# =====================================================================
add_heading_styled("The 3 Things That Cost You the Most Points", level=1)

# --- 1. Comma Splices ---
add_heading_styled("1. Comma Splices (Q6, Q7, Q9, and Q11)", level=2)

p = doc.add_paragraph("A ")
bold(p, "comma splice")
p.add_run(
    " is when you stick two complete sentences together with just a "
    "comma \u2014 but no connecting word. This happened a lot on your test."
)

# Table
t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
t.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header(t, ["Question", "What you wrote", "The problem"])

splice_rows = [
    (
        "Q6",
        "\u201c...popular with fans, it is a fun event and it sells out "
        "every year, plus the fans look forward to it every year.\u201d",
        "Three separate thoughts joined only by commas",
    ),
    (
        "Q7",
        "\u201cOne thing I would change about the event is, the fans "
        "should have to clean up...\u201d",
        "Comma after \u201cis\u201d sticks two sentences together",
    ),
    (
        "Q9",
        "\u201c...it is a good way to help others, Because all the "
        "Teddy Bears...\u201d",
        "Comma before \u201cBecause\u201d (also, \u201cBecause\u201d "
        "shouldn\u2019t be capitalized mid-sentence)",
    ),
    (
        "Q11",
        "\u201c...affected the people of Garissa for the better, it "
        "helped people learn things...\u201d",
        "Two complete thoughts joined by just a comma",
    ),
]
for row_data in splice_rows:
    table_add(t, row_data)

for row in t.rows:
    row.cells[0].width = Inches(0.7)
    row.cells[1].width = Inches(3.3)
    row.cells[2].width = Inches(2.5)

doc.add_paragraph()

p = doc.add_paragraph()
bold(p, "How to catch comma splices \u2014 the Finger Test:")

p = doc.add_paragraph(
    "Every time you use a comma, put your finger over it and read "
    "what\u2019s on the left side. Ask yourself: "
)
italic(p, "\u201cIs this a complete sentence on its own?\u201d")
p.add_run(" Then read the right side. ")
italic(p, "\u201cIs this a complete sentence too?\u201d")

p = doc.add_paragraph("If ")
bold(p, "both sides")
p.add_run(
    " are complete sentences, you have a comma splice! To fix it, "
    "add a connecting word after the comma \u2014 like "
)
bold(p, "and")
p.add_run(", ")
bold(p, "because")
p.add_run(", ")
bold(p, "so")
p.add_run(", or ")
bold(p, "but")
p.add_run(".")

p = doc.add_paragraph()
bold(p, "Here\u2019s what that looks like for Q6:")

p = doc.add_paragraph(style="List Bullet")
p.add_run("Your version: ")
italic(
    p,
    "\u201cThis is one reason why I think the Teddy Bear Toss is "
    "popular with fans, it is a fun event...\u201d",
)

p = doc.add_paragraph(style="List Bullet")
p.add_run("Fixed: ")
italic(p, "\u201cI think the Teddy Bear Toss is popular with fans ")
bold_italic(p, "because")
italic(p, " it is a fun event that sells out every year.\u201d")

p = doc.add_paragraph(
    "See how adding \u201cbecause\u201d connects the two ideas and "
    "makes the sentence stronger?"
)

# --- 2. Circular Reasoning ---
add_heading_styled("2. Circular Reasoning (Q6 and Q9)", level=2)

p = doc.add_paragraph(
    "Circular reasoning is when your \u201creason\u201d just says the "
    "same thing again using different words. Your answer goes in a "
    "circle instead of moving forward."
)

p = doc.add_paragraph()
bold(p, "Q9: ")
italic(
    p,
    "\u201cYes, I do think it is a good way to help others, Because "
    "all the Teddy Bears on the ice after the game go to charity "
    "which is a good way to help others.\u201d",
)

p = doc.add_paragraph("Look at how this starts and ends:")

p = doc.add_paragraph(style="List Bullet")
p.add_run("Start: \u201cgood way to help others\u201d")
p = doc.add_paragraph(style="List Bullet")
p.add_run("End: \u201cgood way to help others\u201d")

p = doc.add_paragraph(
    "Your reason didn\u2019t add anything new \u2014 it just repeated "
    "your answer!"
)

p = doc.add_paragraph()
bold(
    p,
    "How to fix this \u2014 ask yourself: "
    "\u201cAm I telling the reader something they didn\u2019t already know?\u201d",
)

p = doc.add_paragraph(style="List Bullet")
p.add_run("Circular: ")
italic(p, "\u201cIt helps others because it\u2019s a good way to help.\u201d")

p = doc.add_paragraph(style="List Bullet")
p.add_run("Not circular: ")
italic(
    p,
    "\u201cIt helps others because the teddy bears are donated to "
    "charities that give them to children during the holidays.\u201d",
)

p = doc.add_paragraph("The second version adds ")
bold(p, "specific details from the passage")
p.add_run(
    " \u2014 donated, charities, children, holidays. "
    "That\u2019s new information!"
)

p = doc.add_paragraph()
bold(p, "Easy way to remember: ")
p.add_run(
    "After you write your reason, cover up the first half of your "
    "sentence. If the second half says almost the same thing as the "
    "first half, go back to the passage and find a specific fact to "
    "use instead."
)

# --- 3. Reading the Task ---
add_heading_styled("3. Q1 \u2014 Reading the Task Carefully", level=2)

p = doc.add_paragraph()
bold(p, "Task: ")
italic(
    p,
    "\u201cRewrite this sentence to make it clear and concise: That day, "
    "the Hershey Bears broke their own record by collecting 45,650 "
    "stuffed animals that day.\u201d",
)

p = doc.add_paragraph("The key word was ")
bold(p, "concise")
p.add_run(
    ", which means \u201ctake out what you don\u2019t need.\u201d "
    "The sentence says \u201cthat day\u201d twice \u2014 your job was "
    "to remove the extra one."
)

p = doc.add_paragraph(style="List Bullet")
p.add_run("Your version: ")
italic(
    p,
    "\u201cThat day, the Hershey Bears broke their own record, by "
    "collecting 45,650 stuffed animals that day.\u201d",
)
p.add_run(
    " \u2014 You added a comma, but both \u201cthat day\u201ds are still there."
)

p = doc.add_paragraph(style="List Bullet")
p.add_run("Fixed: ")
italic(
    p,
    "\u201cThat day, the Hershey Bears broke their own record by "
    "collecting 45,650 stuffed animals.\u201d",
)

p = doc.add_paragraph()
bold(p, "Strategy: ")
p.add_run(
    "Before you start writing your answer for Q1\u2013Q5, read the "
    "question slowly and ask yourself: "
)
italic(
    p,
    "\u201cWhat exactly is wrong with this sentence? What am I "
    "supposed to fix?\u201d",
)
p.add_run(
    " Then after you write your answer, re-read the question to make "
    "sure you actually did what it asked."
)

# =====================================================================
# Q11 PARAGRAPH
# =====================================================================
add_heading_styled(
    "Q11 Paragraph \u2014 Your Biggest Chance to Earn More Points", level=1
)

p = doc.add_paragraph("Here\u2019s what you wrote:")

quote = doc.add_paragraph()
quote.paragraph_format.left_indent = Inches(0.5)
quote.paragraph_format.right_indent = Inches(0.5)
run = quote.add_run(
    "The Camel Mobile Library affected the people of Garissa for the "
    "better, it helped people learn things that they might really need "
    "in the future, and facts that they can tell their parents and "
    "start a conversation. These books can also be used by adults or "
    "parents to tell their kids stories. It has helped generations in "
    "many different ways."
)
run.italic = True

p = doc.add_paragraph(
    "You had the right idea \u2014 you knew the library helped people "
    "learn, and you thought about both adults and children. But this "
    "paragraph scored 14/20 because of three things:"
)

issues = [
    (
        "No specific evidence from the article",
        " \u2014 \u201cthings that they might really need\u201d is too "
        "vague. What things exactly? The rubric says \u201cuse evidence "
        "from the selection.\u201d",
    ),
    (
        "Details without explanations",
        " \u2014 Even when you mention an idea (like parents telling "
        "stories), you don\u2019t explain why that matters or how the "
        "library made it possible.",
    ),
    (
        "The ending",
        " (\u201cIt has helped generations in many different ways\u201d) "
        "just repeats the beginning without adding anything new.",
    ),
]
for i, (b, rest) in enumerate(issues, 1):
    p = doc.add_paragraph(f"{i}. ")
    bold(p, b)
    p.add_run(rest)

# SPO
add_heading_styled("Use an SPO (Single Paragraph Outline) Before You Write", level=2)

p = doc.add_paragraph(
    "Before you start writing your paragraph, spend 2\u20133 minutes "
    "making a quick plan. Write down:"
)

spo_items = [
    (
        "T.S. (Topic Sentence):",
        " Answer the question from the prompt. Name the topic + your main idea.",
    ),
    (
        "Detail 1 + Explanation:",
        " One specific fact from the article, then explain why it matters "
        "or how it connects to your main idea.",
    ),
    (
        "Detail 2 + Explanation:",
        " Another specific fact from the article, then explain why it "
        "matters or how it connects to your main idea.",
    ),
    (
        "C.S. (Concluding Sentence):",
        " Say your main idea again, but in different words. Connect it "
        "to the bigger picture.",
    ),
]
for b, rest in spo_items:
    p = doc.add_paragraph(style="List Bullet")
    bold(p, b)
    p.add_run(rest)

p = doc.add_paragraph("Here\u2019s what an SPO for this prompt could look like:")

spo_t = doc.add_table(rows=6, cols=2)
spo_t.style = "Table Grid"
spo_t.alignment = WD_TABLE_ALIGNMENT.CENTER

spo_data = [
    (
        "T.S.",
        "The Camel Mobile Library has helped the people of Garissa by "
        "bringing books and learning to a place with no libraries.",
    ),
    (
        "Detail 1",
        "Children in Garissa can now borrow books from the camels and "
        "learn to read.",
    ),
    (
        "Explanation",
        "This is important because before the library came, many children "
        "had never held a book before.",
    ),
    (
        "Detail 2",
        "The library also helps adults learn new things, like how to take "
        "care of their animals or grow food.",
    ),
    (
        "Explanation",
        "This matters because these skills help families earn money and "
        "live better lives.",
    ),
    (
        "C.S.",
        "Thanks to the Camel Mobile Library, people in Garissa now have "
        "a chance to learn and build a better future.",
    ),
]

for i, (label, content) in enumerate(spo_data):
    cl = spo_t.rows[i].cells[0]
    cr = spo_t.rows[i].cells[1]
    cl.text = ""
    cr.text = ""
    run = cl.paragraphs[0].add_run(label)
    run.bold = True
    run.font.size = Pt(10)
    shade_cell(cl, "E8F5E9")
    run2 = cr.paragraphs[0].add_run(content)
    run2.font.size = Pt(10)
    cl.width = Inches(1.0)
    cr.width = Inches(5.5)

doc.add_paragraph()

p = doc.add_paragraph()
bold(p, "The key rule: ")
p.add_run("Every detail needs two parts \u2014 the ")
bold(p, "fact")
p.add_run(" (what happened) and the ")
bold(p, "explanation")
p.add_run(
    " (why it matters or how it connects to your main idea). A detail "
    "without an explanation is like giving someone directions but leaving "
    "out the street names \u2014 they won\u2019t know where you\u2019re going."
)

# =====================================================================
# TEST-TAKING STRATEGIES
# =====================================================================
add_heading_styled("Test-Taking Strategies for Next Time", level=1)

add_heading_styled("Before you write:", level=2)

strats = [
    (
        "Read the rubric box on the test",
        " \u2014 it tells you exactly how you\u2019re scored (writing "
        "skill, grammar, spelling, punctuation). Think of it as your checklist.",
    ),
    (
        "For Q1\u2013Q5,",
        " read the question slowly and figure out what specific skill "
        "it\u2019s asking for (combine, rewrite, fix, make concise). After "
        "you write your answer, re-read the question to make sure you "
        "actually did what it asked.",
    ),
    (
        "For Q11,",
        " build your SPO first. Don\u2019t start writing the paragraph "
        "until your outline has a T.S., 2 details with explanations, and a C.S.",
    ),
]
for i, (b, rest) in enumerate(strats, 1):
    p = doc.add_paragraph(f"{i}. ")
    bold(p, b)
    p.add_run(rest)

add_heading_styled("After you write \u2014 3-Step Proofread:", level=2)

pt = doc.add_table(rows=1, cols=2)
pt.style = "Table Grid"
pt.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header(pt, ["Check", "What to look for"])

proof = [
    (
        "Finger Test",
        "Find every comma. Is there a complete sentence on both sides? "
        "If yes, did I add a connecting word (and, because, so, but)?",
    ),
    (
        "New Info Check",
        "Cover the first half of my sentence. Does the second half tell "
        "the reader something new?",
    ),
    (
        "Bookends Check",
        "Does my sentence start with a capital letter and end with a "
        "period or question mark?",
    ),
]
for check, what in proof:
    row = pt.add_row()
    row.cells[0].text = ""
    row.cells[1].text = ""
    run = row.cells[0].paragraphs[0].add_run(check)
    run.bold = True
    run.font.size = Pt(10)
    run2 = row.cells[1].paragraphs[0].add_run(what)
    run2.font.size = Pt(10)

for row in pt.rows:
    row.cells[0].width = Inches(1.3)
    row.cells[1].width = Inches(5.2)

doc.add_paragraph()

add_heading_styled("Time management:", level=2)
p = doc.add_paragraph("You left ")
bold(p, "Q10 blank")
p.add_run(
    " \u2014 that\u2019s 2 points you could have earned. If you\u2019re "
    "running low on time, the sentence-combining questions (like Q10) "
    "are usually the fastest to answer. Write something, even if it\u2019s "
    "not perfect \u2014 a partial answer can still earn points, but a blank "
    "is always zero."
)

# =====================================================================
# SUMMARY
# =====================================================================
add_heading_styled("Summary \u2014 Your 3 Focus Areas", level=1)

st = doc.add_table(rows=1, cols=3)
st.style = "Table Grid"
st.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header(st, ["Priority", "Skill", "What to practice"])

summary = [
    ("1", "Fix comma splices", "Use the Finger Test on every comma before you submit"),
    (
        "2",
        "Give real reasons, not circular ones",
        "Ask: \u201cDoes my reason tell the reader something new?\u201d "
        "Use specific facts from the passage.",
    ),
    (
        "3",
        "Plan your Q11 paragraph",
        "Always build an SPO first \u2014 T.S., Detail + Explanation, "
        "Detail + Explanation, C.S.",
    ),
]
for priority, skill, practice in summary:
    row = st.add_row()
    for j, text in enumerate([priority, skill, practice]):
        c = row.cells[j]
        c.text = ""
        run = c.paragraphs[0].add_run(text)
        run.font.size = Pt(10)
        if j == 0:
            run.bold = True
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for row in st.rows:
    row.cells[0].width = Inches(0.7)
    row.cells[1].width = Inches(2.0)
    row.cells[2].width = Inches(3.8)

doc.add_paragraph()

p = doc.add_paragraph(
    "You have great instincts, Zara \u2014 you understand what the "
    "questions are asking, and you always try to give reasons. The next "
    "step is catching those comma splices and making your evidence more "
    "specific. If you focus on those two things, you\u2019re going to see "
    "a big jump in your scores!"
)

# Save
out_path = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "Reports",
    "Zara_Lapointe_G3.2_Coaching_Report.docx",
)
os.makedirs(os.path.dirname(out_path), exist_ok=True)
doc.save(out_path)
print(f"Saved to {out_path}")
